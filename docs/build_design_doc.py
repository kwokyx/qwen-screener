"""生成学年设计报告：《基于千问 Agent 的 A 股智能选股系统》。

文档结构仿照课程设计报告模板：
- 评审意见
- 学校封面
- 内容提要、关键词、参考书目
- 目录
- 1 概述
- 2 基本概念和方法、工具
- 3 设计方案
- 4 源代码及实现
- 5 使用说明
- 6 系统测试
- 7 总结与展望

输出：
docs/基于千问的股票筛选系统设计与实现.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SONG = "宋体"
HEI = "黑体"
KAI = "楷体"
EN = "Times New Roman"
MONO = "Consolas"


def set_run_font(
    run,
    *,
    cn: str = SONG,
    en: str = EN,
    size: float = 12,
    bold: bool = False,
    color: str | None = None,
):
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)


def set_paragraph_format(
    p,
    *,
    indent: bool = True,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
    first_indent_pt: float | None = None,
):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent_pt is not None:
        pf.first_line_indent = Pt(first_indent_pt)
    elif indent:
        pf.first_line_indent = Pt(24)
    else:
        pf.first_line_indent = Pt(0)
    return p


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    indent: bool = True,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size: float = 12,
    bold: bool = False,
    cn: str = SONG,
    en: str = EN,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        indent=indent,
        align=align,
        line_spacing=line_spacing,
        space_before=space_before,
        space_after=space_after,
    )
    if text:
        run = p.add_run(text)
        set_run_font(run, cn=cn, en=en, size=size, bold=bold)
    return p


def add_run_line(doc: Document, label: str, value: str, *, size: float = 12):
    p = add_paragraph(doc, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=size, space_after=4)
    r = p.add_run(label)
    set_run_font(r, cn=HEI, en=EN, size=size, bold=True)
    r = p.add_run(value)
    set_run_font(r, cn=SONG, en=EN, size=size)
    return p


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        indent=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_before=12,
        space_after=12,
    )
    run = p.add_run(text)
    set_run_font(run, cn=HEI, en=EN, size=16, bold=True)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        indent=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        space_before=10,
        space_after=6,
    )
    run = p.add_run(text)
    set_run_font(run, cn=HEI, en=EN, size=15, bold=True)
    return p


def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        indent=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        space_before=8,
        space_after=4,
    )
    run = p.add_run(text)
    set_run_font(run, cn=HEI, en=EN, size=14, bold=True)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        indent=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_before=4,
        space_after=8,
    )
    run = p.add_run(text)
    set_run_font(run, cn=SONG, en=EN, size=10.5, bold=True)
    return p


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    insert_before = {
        qn("w:shd"),
        qn("w:noWrap"),
        qn("w:tcMar"),
        qn("w:textDirection"),
        qn("w:tcFitText"),
        qn("w:vAlign"),
        qn("w:hideMark"),
    }
    for idx, child in enumerate(list(tcPr)):
        if child.tag in insert_before:
            tcPr.insert(idx, tcBorders)
            return
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    insert_before = {
        qn("w:noWrap"),
        qn("w:tcMar"),
        qn("w:textDirection"),
        qn("w:tcFitText"),
        qn("w:vAlign"),
        qn("w:hideMark"),
    }
    for idx, child in enumerate(list(tcPr)):
        if child.tag in insert_before:
            tcPr.insert(idx, shd)
            return
    tcPr.append(shd)


def add_table(
    doc: Document,
    header: list[str],
    rows: list[list[str]],
    *,
    col_widths: list[float] | None = None,
    caption: str | None = None,
):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

    for i, title in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_borders(cell)
        set_cell_shading(cell, "DCE6F1")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(title)
        set_run_font(run, cn=HEI, en=EN, size=10.5, bold=True)

    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.rows[r].cells[c]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(value)
            set_run_font(run, cn=SONG, en=EN, size=10.5)
    doc.add_paragraph()
    return table


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        indent=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.15,
        space_before=4,
        space_after=8,
    )
    p.paragraph_format.left_indent = Pt(12)
    run = p.add_run(code)
    set_run_font(run, cn=SONG, en=MONO, size=9.5)
    return p


def add_numbered_text(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph()
        set_paragraph_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        r = p.add_run(item)
        set_run_font(r, cn=SONG, en=EN, size=12)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style = doc.styles["Normal"]
    style.font.name = EN
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), SONG)
    rFonts.set(qn("w:ascii"), EN)
    rFonts.set(qn("w:hAnsi"), EN)

    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(qn("w:percent"), "100")
    if qn("w:val") in zoom.attrib:
        del zoom.attrib[qn("w:val")]


def add_review_page(doc: Document):
    add_paragraph(doc, "评审意见：", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=14, bold=True, cn=HEI)
    for _ in range(18):
        add_paragraph(doc, "", indent=False)
    add_paragraph(doc, "指导教师（签字）：", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, bold=True)
    add_paragraph(doc, "年    月    日", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
    add_page_break(doc)


def add_cover(doc: Document):
    for text, size, bold, space in [
        ("西南大学", 18, True, 8),
        ("计算机与信息科学学院", 18, True, 36),
        ("学年设计报告", 26, True, 54),
    ]:
        p = add_paragraph(
            doc,
            text,
            indent=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=size,
            bold=bold,
            cn=HEI,
            space_after=space,
        )
        p.paragraph_format.line_spacing = 1.4

    cover_rows = [
        ("课　　程", "数据科学行业应用综合实践（学年设计）"),
        ("题　　目", "基于千问 Agent 的 A 股智能选股系统设计与实现"),
        ("级、专业", "2023级 计算机科学与技术（中外合作） 计科中外班"),
        ("学生姓名", "成员1（组长）、成员2、成员3、成员4、成员5"),
        ("提交日期", "2026年6月12日"),
    ]
    for label, value in cover_rows:
        p = add_paragraph(doc, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=12)
        r = p.add_run(f"{label}：")
        set_run_font(r, cn=HEI, en=EN, size=14, bold=True)
        r = p.add_run(value)
        set_run_font(r, cn=SONG, en=EN, size=14)
    add_page_break(doc)


def add_front_matter(doc: Document):
    add_caption(doc, "内容提要：")
    add_paragraph(
        doc,
        "本报告围绕“基于千问 Agent 的 A 股智能选股系统设计与实现”展开。"
        "系统面向个人投资者和课程演示场景，提供登录注册、行情概览、AI 自然语言选股、条件选股、策略选股、"
        "自选股管理、价格预警、个股详情、K 线展示、千问解读、数据健康检查和飞书推送等功能。"
        "项目采用前后端分离架构，后端基于 FastAPI、SQLAlchemy、APScheduler 和本地股票数据库实现业务接口、"
        "筛选引擎、定时同步和通知推送；前端采用 Vue 3、Vite、Pinia、Naive UI 与 klinecharts 构建桌面端交互界面。"
    )
    add_paragraph(
        doc,
        "系统的核心特点是把大语言模型限制在“理解意图和选择工具”的边界内。AI 选股使用 bounded ReAct："
        "模型只输出普通回复 final 或白名单工具 action，后端负责参数校验、字段能力边界检查、工具执行和结果汇总。"
        "当 AI 服务超时或不可达时，系统不会把不可靠的本地兜底伪装成筛选结果，而是明确降级提示。"
        "股票筛选和策略计算均由本地确定性引擎执行，支持 PE、PB、市值、股息率、ROE、营收同比、净利润同比、"
        "毛利率、负债率、行业、板块以及部分均线/放量/突破类技术字段；三年 CAGR、扣非净利润、现金流、EPS、"
        "机构持仓和研报评级等未入库字段明确标记为不支持。"
    )
    add_paragraph(
        doc,
        "报告从需求分析、总体架构、数据库设计、Agent 设计、策略引擎、数据同步、前端页面、部署运行和测试验证等方面"
        "说明系统实现过程。系统已经形成可演示的完整闭环：用户可通过自然语言表达选股目标，后端将其转化为结构化工具调用；"
        "也可直接使用条件选股和 6 个内置日线策略；自选页面支持批量编辑、排序和预警规则；数据健康页展示最新交易日、"
        "应至交易日、覆盖率和同步异常，避免把数据新鲜度伪装成正常。"
    )
    add_paragraph(
        doc,
        "关键词：大语言模型；AI Agent；A 股选股；FastAPI；Vue 3；数据同步；策略选股",
        indent=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        bold=True,
    )
    add_paragraph(
        doc,
        "参考书目：\n《软件工程》；《面向对象分析与设计》；FastAPI Documentation；Vue.js Guide；"
        "SQLAlchemy Documentation；APScheduler Documentation；Baostock 帮助文档；AKShare 文档；"
        "通义千问 DashScope API 文档。参见正文中“参考文献”。",
        indent=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_page_break(doc)


def add_manual_toc(doc: Document):
    add_h1(doc, "目　录")
    toc = [
        "1 概述",
        "1.1 问题描述",
        "1.2 设计目标",
        "1.3 任务与成员分工",
        "2 基本概念和方法、工具",
        "2.1 软件开发基本概念",
        "2.2 B/S架构与前后端分离",
        "2.3 关键技术",
        "2.4 运行环境与硬件要求",
        "3 设计方案",
        "3.1 需求分析",
        "3.2 总体架构设计",
        "3.3 功能模块设计",
        "3.4 数据库设计",
        "3.5 Agent 与工具调用设计",
        "3.6 策略与预警设计",
        "3.7 数据同步与健康检查设计",
        "4 源代码及实现",
        "4.1 工程组织与入口实现",
        "4.2 后端接口与业务服务实现",
        "4.3 筛选引擎与字段边界实现",
        "4.4 AI Agent 与千问解读实现",
        "4.5 策略选股、通知与自选实现",
        "4.6 数据同步、配置与部署实现",
        "5 使用说明",
        "5.1 系统启动与入口",
        "5.2 AI 选股使用说明",
        "5.3 条件选股与策略选股使用说明",
        "5.4 自选股、预警和详情页使用说明",
        "6 系统测试",
        "6.1 测试环境与测试方法",
        "6.2 功能测试",
        "6.3 边界条件测试",
        "6.4 性能与数据健康测试",
        "7 总结与展望",
        "参考文献",
    ]
    for item in toc:
        p = add_paragraph(doc, item, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
        p.paragraph_format.left_indent = Pt(24 if item[0].isdigit() and "." in item[:4] else 0)
    add_page_break(doc)


def add_chapter_1(doc: Document):
    add_h1(doc, "1 概述")
    add_h2(doc, "1.1 问题描述")
    add_paragraph(
        doc,
        "A 股市场股票数量多、指标来源分散，普通用户在进行选股时通常需要同时理解估值、盈利能力、成长性、"
        "行业分类、K 线形态和风险警示等概念。传统条件筛选工具虽然能力完整，但需要用户手动选择字段、运算符和阈值；"
        "纯自然语言问答虽然表达方便，但如果模型直接生成股票名单，容易出现字段幻觉、数据口径不一致和结果不可复现等问题。"
    )
    add_paragraph(
        doc,
        "因此，本项目把“自然语言理解”和“确定性数据筛选”分开设计：用户可以用一句话说明目标，例如"
        "“PE 低于 15、ROE 大于 15 的银行股”或“找 RPS 强势突破的股票”；AI Agent 只负责判断这是筛选、策略、"
        "详情、解释、排序、分页还是普通对话；真正的数据查询由后端本地筛选引擎和策略引擎完成。"
        "这种设计既降低了使用门槛，也保证结果可校验、可复现。"
    )
    add_paragraph(
        doc,
        "项目还需要解决数据时效性问题。股票行情和财务数据来自外部数据源，网络延迟、接口限流、交易日差异和补充字段缺失"
        "都会影响用户判断。系统通过定时任务、手动同步、数据健康检查和同步异常提示来展示“最新交易日、应至交易日、"
        "覆盖率、任务状态和下一步操作”，避免在数据未完成时误导用户。"
    )

    add_h2(doc, "1.2 设计目标")
    add_paragraph(
        doc,
        "本项目的总体目标是设计并实现一个可运行、可演示、可验证的 A 股智能选股系统。系统既要提供传统条件选股和"
        "策略选股能力，也要提供自然语言 AI 选股入口，并在数据边界、模型降级和用户交互上保持清晰。"
    )
    add_numbered_text(
        doc,
        [
            "1. 功能目标：完成登录注册、行情概览、AI 选股、条件选股、策略选股、自选股、预警通知、个股详情和千问解读。",
            "2. 数据目标：本地保存股票基础信息、日线行情、财务指标、分红记录、自选股、通知和对话历史，支持定时同步。",
            "3. Agent 目标：采用 bounded ReAct，不实现无限自主 Agent；模型必须选择白名单工具或返回普通回复。",
            "4. 安全目标：不在代码和文档中保存真实密钥；密码哈希存储；筛选字段和工具参数由后端 schema 强校验。",
            "5. 演示目标：Docker Compose 与 Railway 均可部署，普通测试不依赖真实 AI，真实 AI 仅在 smoke 阶段验证。",
        ],
    )

    add_h2(doc, "1.3 任务与成员分工")
    add_paragraph(
        doc,
        "本项目按后端、前端、数据库、数据采集与维护、数据处理与质量保障五个方向划分工作。各方向在实际联调时相互配合，"
        "但公开分工保持如下表所示。"
    )
    add_table(
        doc,
        header=["成员", "角色", "主要工作内容"],
        rows=[
            [
                "成员1",
                "后端开发与 AI Agent",
                "负责后端接口、核心业务逻辑和 AI 自然语言选股 Agent 开发，完成股票筛选、策略选股、自选股和预警规则等功能的服务端实现。",
            ],
            [
                "成员2",
                "前端开发",
                "负责登录页、AI 选股页、市场概览页、条件选股页、策略选股页、自选股页和股票详情页等页面开发，实现数据展示和用户交互。",
            ],
            [
                "成员3",
                "数据库设计",
                "负责股票基础信息、行情数据、财务指标、历史 K 线、自选股和预警规则等数据表设计，并维护字段和表结构。",
            ],
            [
                "成员4",
                "数据采集与维护",
                "负责股票行情、指数、板块、财务指标和历史 K 线等数据来源整理、采集规则维护和同步状态记录。",
            ],
            [
                "成员5",
                "数据处理与质量保障",
                "负责对采集到的数据进行格式统一、字段校验、缺失值和异常值处理，检查数据覆盖率和更新时间，并参与主要功能测试和演示前验收。",
            ],
        ],
        col_widths=[1.8, 3.8, 10.0],
        caption="表1-1 小组成员分工",
    )


def add_chapter_2(doc: Document):
    add_h1(doc, "2 基本概念和方法、工具")
    add_h2(doc, "2.1 软件开发基本概念")
    add_paragraph(
        doc,
        "软件工程强调需求、设计、实现、测试和维护之间的完整过程。本系统采用分层设计方法，将前端界面、API 路由、"
        "业务服务、数据模型和外部依赖分别封装，减少模块之间的耦合。股票、用户、自选股、通知、对话历史和同步任务"
        "被抽象为相对稳定的数据对象；筛选、策略、AI 工具调用和数据同步则作为业务服务组织。"
    )
    add_paragraph(
        doc,
        "系统实现时遵循“能力边界优先”的原则。金融数据字段并不是模型想使用就能使用，只有已经落入本地数据库、"
        "有明确数据来源、能通过后端校验并能被 SQL 或策略计算执行的字段，才允许参与筛选。缺少数据表或口径不稳定的字段，"
        "例如三年净利润复合增速、经营现金流、EPS 和机构持仓，系统会明确提示不支持。"
    )

    add_h2(doc, "2.2 B/S架构与前后端分离")
    add_paragraph(
        doc,
        "系统采用浏览器/服务器模式。用户通过浏览器访问 Vue 单页应用，前端通过 HTTP 和 SSE 调用后端 API；"
        "后端负责认证、业务规则、AI 调用、工具执行、数据库访问和定时任务；数据库保存结构化业务数据。"
        "前后端分离使页面开发和接口开发可以并行进行，也便于 Docker 和 Railway 分别部署。"
    )
    add_paragraph(
        doc,
        "在 AI 选股场景中，前后端分离还承担了重要的安全边界。前端只展示用户输入、模型回复、工具调用轨迹和结果列表；"
        "后端才有权决定是否执行筛选、策略、详情查询或飞书推送。这样可以避免前端把模型文本误当作真实筛选结果。"
    )

    add_h2(doc, "2.3 关键技术")
    add_table(
        doc,
        header=["技术或工具", "用途", "在本系统中的作用"],
        rows=[
            ["FastAPI", "后端应用框架", "提供 REST API、依赖注入、请求校验、Swagger 文档和 SSE 流式接口。"],
            ["SQLAlchemy 2.0", "ORM 数据访问", "定义股票、用户、自选、通知、对话历史等模型，屏蔽 SQLite/MySQL 差异。"],
            ["SQLite / MySQL", "关系型数据库", "保存基础信息、行情、财务、分红、策略上下文、自选和通知数据。"],
            ["Redis", "可选缓存", "缓存 AI 解析、个股分析和部分行情聚合，连接失败时业务自动降级。"],
            ["APScheduler", "定时任务", "实现收盘后行情同步、周度财务/分红同步、策略推送和数据库备份。"],
            ["Baostock", "主数据源", "同步 A 股基础信息、日/周/月 K、分钟 K、财务和分红数据。"],
            ["AKShare / Sina", "补充数据源", "提供北交所、实时行情和少量兼容兜底数据。"],
            ["OpenAI SDK / DashScope", "AI 调用", "以 OpenAI 兼容接口或 DashScope 调用模型，完成 Agent action/final 和千问解读。"],
            ["Vue 3 + Vite", "前端框架", "构建单页应用、路由、页面组件和开发热更新。"],
            ["Pinia + Vue Router", "状态与路由", "维护登录态、自选状态、页面切换和详情跳转。"],
            ["Naive UI", "组件库", "实现表格、表单、弹窗、按钮、通知中心和页面布局。"],
            ["klinecharts", "K 线图", "在个股详情页展示日/周/月 K 线和成交量。"],
            ["Docker Compose", "本地部署", "统一启动 backend、frontend 和 redis，便于演示复现。"],
            ["Railway", "线上部署", "前端和后端分别作为服务部署，环境变量在平台侧配置。"],
        ],
        col_widths=[3.2, 2.8, 9.4],
        caption="表2-1 系统关键技术",
    )

    add_h2(doc, "2.4 运行环境与硬件要求")
    add_paragraph(
        doc,
        "开发和演示环境可以使用普通个人电脑。后端要求 Python 3.10 及以上版本，前端要求 Node.js 与 npm，"
        "容器部署需要 Docker 和 Docker Compose。数据库默认使用 SQLite，若部署到生产环境可切换为 MySQL；"
        "Redis 为可选缓存，未配置时不会影响核心功能。浏览器建议使用 Chrome、Edge 或 Safari。"
    )
    add_table(
        doc,
        header=["环境项", "推荐配置"],
        rows=[
            ["操作系统", "macOS / Windows / Linux 均可，Docker 环境建议 Linux 或 macOS。"],
            ["CPU 与内存", "4 核 CPU、8GB 内存可运行；全量同步和前后端构建建议 16GB 内存。"],
            ["后端运行", "Python 3.10+、FastAPI、Uvicorn、SQLAlchemy、APScheduler。"],
            ["前端运行", "Node.js 20+、Vite、Vue 3、Naive UI。"],
            ["数据存储", "SQLite 适合课程演示，MySQL 8 适合长期部署；Redis 可选。"],
            ["外部配置", "AI Key、飞书应用参数等通过环境变量配置，不写入仓库。"],
        ],
        col_widths=[3.2, 12.0],
        caption="表2-2 运行环境要求",
    )


def add_chapter_3(doc: Document):
    add_h1(doc, "3 设计方案")
    add_h2(doc, "3.1 需求分析")
    add_paragraph(
        doc,
        "系统需求可以按照用户场景分为六类：行情浏览、股票筛选、策略选股、个股分析、自选预警和系统运维。"
        "游客可以查看公开行情、执行筛选和查看详情；登录用户可以保存自选股、设置预警、查看通知和保留对话历史；"
        "系统维护者需要查看数据健康、触发同步、检查 AI 状态和处理异常任务。"
    )
    add_table(
        doc,
        header=["角色", "核心需求", "主要输出"],
        rows=[
            ["游客", "浏览市场、搜索股票、试用 AI 选股、查看详情", "指数、板块、筛选结果、K 线和基础指标"],
            ["登录用户", "保存自选、批量编辑、设置预警、保留对话历史", "自选列表、通知中心、预警规则、历史对话"],
            ["系统维护者", "同步数据、检查数据健康、查看任务异常、配置 AI 和飞书", "健康报告、同步状态、release smoke 结果"],
        ],
        col_widths=[2.6, 6.2, 6.4],
        caption="表3-1 角色需求分析",
    )

    add_h2(doc, "3.2 总体架构设计")
    add_paragraph(
        doc,
        "系统总体上由前端展示层、后端服务层、数据存储层、外部数据源和 AI 服务五部分组成。前端页面通过 /api/v1 前缀访问"
        "后端接口，AI 选股使用 SSE 流式接口返回普通对话、工具调用和结果。后端服务层内部再拆分为 auth、stock、screener、"
        "qwen、strategy、market、chat、notifications 和 health 等 API 模块。"
    )
    add_caption(doc, "图3-1 系统总体架构图")
    add_code_block(
        doc,
        "浏览器 Vue 3 SPA\n"
        "  |-- Dashboard / AI选股 / 条件选股 / 策略选股 / 自选 / 详情 / 登录\n"
        "  |\n"
        "  | HTTP + SSE (/api/v1/*)\n"
        "  v\n"
        "FastAPI 后端\n"
        "  |-- API: auth, stock, screener, qwen, strategy, market, chat, notifications, health\n"
        "  |-- Services: screener_engine, agent_react, strategy_selector, data_sync, scheduler, feishu, cache\n"
        "  |-- Models: users, stock_basic, stock_daily, stock_financial, stock_dividend, watchlist, notifications, chat_sessions\n"
        "  |\n"
        "  | SQLAlchemy ORM / Redis Client / AI Client\n"
        "  v\n"
        "SQLite 或 MySQL + Redis 可选缓存\n"
        "  |\n"
        "  | 定时同步 / 手动同步\n"
        "  v\n"
        "Baostock 主数据源 + AKShare/Sina 兜底 + AI OpenAI-compatible/DashScope + 飞书通知",
    )

    add_h2(doc, "3.3 功能模块设计")
    add_table(
        doc,
        header=["模块", "子功能", "设计说明"],
        rows=[
            ["认证模块", "验证码、注册、登录、当前用户", "登录和注册均带一次性图形验证码；密码哈希存储；JWT 用于鉴权。"],
            ["行情概览", "指数、板块、涨跌榜、ticker、数据新鲜度", "Dashboard 首屏展示 6 个宽基指数、市场异动和数据健康提示。"],
            ["AI 选股", "普通对话、筛选、解释、排序、分页、详情、策略", "bounded ReAct 由模型选择白名单工具；后端执行并返回结构化结果。"],
            ["条件选股", "字段、操作符、阈值、排序、保存策略", "用户通过表单直接组合条件，结果可批量加入自选。"],
            ["策略选股", "6 个内置日线策略", "本地读取最近 N 日全市场 K 线，按策略规则计算命中强度。"],
            ["个股详情", "基础信息、行情、K 线、财务、同行业、千问解读", "实时行情失败时仍显示本地详情和 K 线。"],
            ["自选与预警", "自选同步、排序、批量编辑、价格/涨跌幅预警", "预警触发后写通知中心，并可调用后端飞书推送。"],
            ["数据健康", "latest/expected/fresh/coverage/sync_warnings", "明确展示数据是否达到最近应有交易日，不伪造 fresh=true。"],
        ],
        col_widths=[2.8, 4.4, 8.2],
        caption="表3-2 功能模块设计",
    )

    add_h2(doc, "3.4 数据库设计")
    add_paragraph(
        doc,
        "数据库设计以股票时序数据和用户业务数据为中心。股票基础信息和时序行情分表保存，避免行情表重复保存股票名称和行业；"
        "用户、自选、通知、对话历史等表独立保存，便于登录态与个人数据扩展。"
    )
    add_table(
        doc,
        header=["表名", "中文含义", "关键字段", "用途"],
        rows=[
            ["users", "用户表", "id, username, email, hashed_password, is_active", "保存登录账号和密码哈希。"],
            ["stock_basic", "股票基础信息", "code, name, industry, market, list_date, total_share", "保存全 A 股票名称、行业、板块和上市信息。"],
            ["stock_daily", "日线行情与估值", "code, trade_date, open/high/low/close, volume, amount, pe, pb, market_cap, turnover, dividend_yield", "保存行情、估值、市值、换手率和股息率。"],
            ["stock_financial", "财务指标", "code, report_date, roe, net_profit, revenue, revenue_yoy, profit_yoy, gross_margin, debt_ratio", "保存最新财务指标，用于基本面筛选。"],
            ["stock_dividend", "现金分红", "code, operate_date, cash_per_share, notice_date, pay_date", "用于本地重算 TTM 股息率。"],
            ["watchlist", "自选股", "user_id, code, note, alerts, ref_price, created_at", "保存用户自选、加入价和预警规则。"],
            ["notifications", "通知表", "user_id, kind, tone, stock_code, stock_name, title, desc, fired_at", "保存系统通知和价格预警触发记录。"],
            ["chat_sessions", "对话历史", "query, parsed_conditions, items, total, agent_plan, tool_trace, tool_calls", "保存 AI 选股快照，避免历史回看时重复调用 AI。"],
            ["sync_meta", "同步状态", "job, status, updated_at, message", "记录数据同步任务状态、异常和完成时间。"],
        ],
        col_widths=[2.8, 2.4, 6.8, 3.4],
        caption="表3-3 数据库表汇总",
    )

    add_h2(doc, "3.5 Agent 与工具调用设计")
    add_paragraph(
        doc,
        "AI 选股不是无限自主 Agent，而是有边界的 ReAct。模型只允许返回两类结果：final 表示普通对话回复；action 表示"
        "选择后端白名单工具。后端接收 action 后，先进行 schema 校验和字段能力边界校验，再执行本地工具。"
        "工具执行结果由后端确定性总结，不再依赖第二次模型总结。"
    )
    add_table(
        doc,
        header=["工具", "用途", "边界控制"],
        rows=[
            ["stock_screen", "把结构化条件交给筛选引擎执行", "字段必须在白名单内，不支持字段直接停止。"],
            ["explain_result", "解释上一轮筛选结果或某只股票", "必须依赖上下文，不重新筛选。"],
            ["sort/page", "对上一轮结果排序或换一批", "只操作当前结果快照，不调用新筛选。"],
            ["stock_detail", "定位股票详情", "按股票名或代码查详情，不触发筛选。"],
            ["strategy_select", "执行内置策略", "strategy_id 必须在显式注册表内。"],
            ["strategy_design", "解释策略思路或生成条件建议", "不伪造回测收益。"],
            ["ask_clarification", "信息不足或字段不支持时追问", "不返回部分筛选结果冒充完整命中。"],
        ],
        col_widths=[3.2, 5.2, 6.8],
        caption="表3-4 Agent 工具设计",
    )

    add_h2(doc, "3.6 策略与预警设计")
    add_paragraph(
        doc,
        "策略选股是本地日线规则引擎，不是 AI 算分系统。策略类统一继承 BaseStrategy，声明 id、name、tag、description、"
        "rules 和 history_days，并在 registry 中显式注册。执行时读取最近 N 个交易日全市场 K 线，过滤数据不足和关键字段缺失的股票，"
        "再按策略内部规则计算命中强度。"
    )
    add_table(
        doc,
        header=["策略", "策略 id", "窗口", "说明"],
        rows=[
            ["海龟突破", "turtle_breakout", "35 日", "关注突破近期高点且成交额满足条件的强势股。"],
            ["均线放量", "ma_volume", "35 日", "关注短均线强于长均线并伴随成交量放大的股票。"],
            ["RPS 强势突破", "rps_breakout", "130 日", "基于 120 日相对强度和突破信号筛选强势股。"],
            ["高位窄幅整理", "high_tight_flag", "55 日", "寻找前期大幅上涨后高位窄幅整理的形态。"],
            ["涨停后承接", "limit_up_shakeout", "3 日", "关注涨停后短期承接和回踩表现。"],
            ["趋势急跌修复", "uptrend_limit_down", "65 日", "寻找趋势中急跌后仍具备修复条件的股票。"],
        ],
        col_widths=[3.0, 3.4, 2.2, 6.8],
        caption="表3-5 内置策略",
    )
    add_paragraph(
        doc,
        "预警分为站内通知和飞书推送两类。自选股页面保存价格突破、价格跌破、累计涨跌幅、日内涨跌幅等规则；"
        "前端预警引擎按固定间隔读取自选股最新价并评估规则，触发后写入通知中心。后端提供飞书推送接口，策略扫描和预警触发"
        "均可发送到飞书，便于演示“系统发现机会后外部通知”的闭环。"
    )

    add_h2(doc, "3.7 数据同步与健康检查设计")
    add_paragraph(
        doc,
        "数据同步由 APScheduler 随 FastAPI 启动注册。系统不假设数据一定同步成功，而是把每个任务状态写入 sync_meta，"
        "在健康检查接口中输出最新交易日、应至交易日、数据覆盖率、是否新鲜、同步异常和正在运行的任务。"
    )
    add_table(
        doc,
        header=["时间", "任务", "内容"],
        rows=[
            ["周一至周五 15:05", "market_refresh", "收盘后快刷日 K 与估值面。"],
            ["周一至周五 15:30", "daily_market", "全市场日 K 补偿。"],
            ["周一至周五 16:00", "daily_value", "估值与股息率补充。"],
            ["周一至周五 18:00", "strategy_push", "全策略扫描并推送飞书。"],
            ["周六 02:00", "weekly_fundamentals", "财务指标同步。"],
            ["周六 03:00", "weekly_dividend", "现金分红与 TTM 股息率。"],
            ["周日 02:00", "weekly_basic", "股票列表、新股和退市更新。"],
            ["周日 03:00", "weekly_kline_backfill", "全市场近期 K 线回填。"],
            ["每 6 小时", "db_backup", "SQLite 冷备份。"],
        ],
        col_widths=[3.4, 4.2, 7.6],
        caption="表3-6 定时任务设计",
    )


def add_chapter_4(doc: Document):
    add_h1(doc, "4 源代码及实现")
    add_h2(doc, "4.1 工程组织与入口实现")
    add_paragraph(
        doc,
        "项目采用 backend 与 frontend 两个主目录。后端按 API、models、schemas、services 和 scripts 分层；前端按 views、"
        "components、stores、services 和 scripts 组织。Docker Compose 负责本地一键启动，Railway 负责线上前后端分服务部署。"
    )
    add_code_block(
        doc,
        "qwen-stock-screener-naive/\n"
        "  backend/\n"
        "    app/api/               # auth, stock, screener, qwen, market, strategy, notifications, health\n"
        "    app/models/            # SQLAlchemy 模型\n"
        "    app/schemas/           # Pydantic 请求和响应模型\n"
        "    app/services/          # 筛选、Agent、策略、同步、缓存、飞书、调度器\n"
        "    scripts/               # 数据同步、release smoke、agent smoke\n"
        "    tests/                 # pytest 回归测试\n"
        "  frontend/\n"
        "    src/views/             # Login、Dashboard、Chat、Results、Strategy、Portfolio、Detail\n"
        "    src/components/        # K线、布局、通知、通用组件\n"
        "    src/services/          # API、预警引擎、缓存辅助\n"
        "    scripts/               # Playwright smoke\n"
        "  docs/                    # API、策略、字段边界、部署和设计文档\n",
    )

    add_h2(doc, "4.2 后端接口与业务服务实现")
    add_table(
        doc,
        header=["接口模块", "文件", "主要功能"],
        rows=[
            ["认证", "backend/app/api/auth.py", "验证码、注册、登录、当前用户信息。"],
            ["股票", "backend/app/api/stock.py", "搜索、详情、K 线、分钟 K、自选 CRUD。"],
            ["筛选", "backend/app/api/screener.py", "结构化筛选、自然语言一次性筛选、SSE 流式筛选。"],
            ["千问", "backend/app/api/qwen.py", "个股 AI 解读，一次性和 SSE 两种返回。"],
            ["策略", "backend/app/api/strategy.py", "策略模板、内置策略执行、自然语言策略入口。"],
            ["市场", "backend/app/api/market.py", "指数、板块、涨跌榜、ticker 和概览聚合。"],
            ["通知", "backend/app/api/notification.py", "通知中心、已读、清空、飞书预警推送。"],
            ["健康", "backend/app/api/health.py", "AI 探活、数据健康、缓存状态、手动同步、备份。"],
        ],
        col_widths=[2.4, 5.8, 7.0],
        caption="表4-1 后端接口模块",
    )

    add_h2(doc, "4.3 筛选引擎与字段边界实现")
    add_paragraph(
        doc,
        "筛选引擎的核心是把 FilterCondition 转为 SQLAlchemy 表达式。支持字段分为数据库字段和派生技术字段两类。"
        "数据库字段直接映射到 stock_daily、stock_financial 或 stock_basic；派生技术字段基于最近 K 线窗口计算。"
        "字符串字段如 industry 支持中文同义词展开，例如“金融”可匹配银行、证券、保险和多元金融。"
    )
    add_table(
        doc,
        header=["字段类别", "支持字段", "数据来源"],
        rows=[
            ["估值行情", "pe, pb, market_cap, close, turnover, dividend_yield", "stock_daily；股息率由 stock_dividend 重算。"],
            ["财务指标", "roe, revenue_yoy, profit_yoy, gross_margin, debt_ratio", "stock_financial 最新报告期。"],
            ["基础属性", "industry, market, risk_flag", "stock_basic；risk_flag 基于名称派生。"],
            ["技术字段", "ma5, ma20, volume_ratio_20, breakout_20, ma5_above_ma20, pct_change_20", "基于 stock_daily 最近窗口计算。"],
            ["暂不支持", "profit_cagr_3y, deducted_profit, operating_cash_flow, eps, ps, institution_holding, research_rating", "本地未入库或口径不稳定，必须停止筛选并说明原因。"],
        ],
        col_widths=[2.4, 6.8, 6.0],
        caption="表4-2 筛选字段能力边界",
    )
    add_code_block(
        doc,
        "def validate_screen_request(req):\n"
        "    if req.sort_by is not None and req.sort_by not in SORT_FIELDS:\n"
        "        raise ValueError('不支持的排序字段')\n"
        "    for cond in req.conditions:\n"
        "        if cond.field in STRING_FIELDS:\n"
        "            # industry / market 只允许非空字符串 eq 或非空字符串数组 in\n"
        "            validate_string_condition(cond)\n"
        "        elif cond.op == 'between':\n"
        "            validate_two_numeric_values(cond.value)\n"
        "        else:\n"
        "            validate_numeric_condition(cond)\n",
    )

    add_h2(doc, "4.4 AI Agent 与千问解读实现")
    add_paragraph(
        doc,
        "Agent 入口位于 agent_react.py，模型规划位于 qwen_client/agent_planner.py。系统 prompt 明确要求模型在白名单工具中选择，"
        "并减少输出字段，避免过长 schema 导致超时。后端会记录 model_ms、tool_ms 和 fallback_reason，用于观察 AI 慢、"
        "网络不可达或模型返回非法 action 的情况。普通对话不显示工具轨迹，只有确实调用工具时才展示工具选择和执行信息。"
    )
    add_caption(doc, "图4-1 AI 选股执行流程")
    add_code_block(
        doc,
        "用户输入\n"
        "  -> unsupported metric 前置检查\n"
        "  -> 调用 bounded ReAct 模型\n"
        "      -> final: 普通回复，不执行工具\n"
        "      -> action: 后端校验 action/tool/schema\n"
        "          -> stock_screen / strategy_select / stock_detail / explain / sort / page\n"
        "          -> 本地工具执行\n"
        "          -> 后端确定性总结\n"
        "  -> SSE 返回阶段事件与最终结果",
    )
    add_paragraph(
        doc,
        "千问解读用于个股详情页。后端根据股票基础信息、最新行情和财务指标组装上下文，请 AI 生成基本面解释；"
        "如果 AI 服务不可用，详情页仍展示本地数据和 K 线，只在解读区域提示服务不可用。"
    )

    add_h2(doc, "4.5 策略选股、通知与自选实现")
    add_paragraph(
        doc,
        "策略实现位于 backend/app/services/strategies/，每个策略一个独立文件。strategy_selector.py 负责读取最近 N 日 K 线、"
        "缓存策略结果、singleflight 防止并发重复计算、按 limit 截断返回以及可选飞书推送。"
    )
    add_code_block(
        doc,
        "run_strategy_selection(strategy_id, limit, notify=True)\n"
        "  -> 检查 strategy_id 是否在 TEMPLATE_MAP\n"
        "  -> 获取最新交易日，生成缓存 key\n"
        "  -> 读取策略 history_days 对应的全市场 K 线\n"
        "  -> 调用策略类 run(df)\n"
        "  -> 按策略内部 score 排序\n"
        "  -> 截断返回，并按需推送飞书",
    )
    add_paragraph(
        doc,
        "自选股使用 watchlist 表保存 code、note、ref_price 和 alerts JSON。通知中心使用 notifications 表保存触发记录。"
        "前端自选页面支持排序、批量编辑、批量移出和预警批量启停；右下角悬浮按钮展示自选入口，通知中心展示站内预警。"
    )

    add_h2(doc, "4.6 数据同步、配置与部署实现")
    add_paragraph(
        doc,
        "数据同步统一通过 scripts.sync_data 子命令触发，并由 scheduler 定时调度。后端配置通过环境变量读取，真实密钥只放在"
        "本地 .env 或 Railway Variables 中，不提交到仓库。Docker Compose 下前端通过 nginx 反向代理 /api；Railway 下前端和"
        "后端分别部署，前端通过线上 API 地址访问后端。"
    )
    add_table(
        doc,
        header=["命令", "作用"],
        rows=[
            ["python -m scripts.sync_data basic", "拉取全 A 股基础信息。"],
            ["python -m scripts.sync_data daily 5", "拉取全市场最近 5 天日 K。"],
            ["python -m scripts.sync_data financial all", "同步财务指标。"],
            ["python -m scripts.sync_data dividend", "同步现金分红并计算股息率。"],
            ["python -m scripts.sync_data full", "依次执行基础信息、日线、财务和分红同步。"],
            ["python scripts/release_smoke.py", "发布前冒烟检查服务、AI、数据和密钥扫描。"],
        ],
        col_widths=[6.2, 9.0],
        caption="表4-3 常用后端命令",
    )


def add_chapter_5(doc: Document):
    add_h1(doc, "5 使用说明")
    add_h2(doc, "5.1 系统启动与入口")
    add_paragraph(
        doc,
        "本地演示推荐使用 Docker Compose。首次启动前复制 backend/.env.example 为 backend/.env，并填写数据库、AI 和飞书相关"
        "环境变量。真实 API Key 不能写入 README、测试文件或提交记录。"
    )
    add_code_block(
        doc,
        "cp backend/.env.example backend/.env\n"
        "docker compose up -d --build\n"
        "docker compose exec -T backend python -m scripts.sync_data full\n"
        "open http://localhost:8080",
    )
    add_table(
        doc,
        header=["页面", "路径", "说明"],
        rows=[
            ["登录注册", "/login", "验证码、注册、登录、退出重登。"],
            ["行情概览", "/dashboard", "指数、市场概况、板块涨跌、市场异动和数据新鲜度。"],
            ["AI 选股", "/chat", "用自然语言筛选 A 股，支持普通对话和工具调用。"],
            ["条件选股", "/results", "用字段、操作符和阈值组合条件筛选。"],
            ["策略选股", "/strategy", "条件策略与 6 个内置日线策略。"],
            ["自选预警", "/portfolio", "自选列表、排序、批量编辑、预警规则和通知中心。"],
            ["个股详情", "/detail/600036.SH", "基础信息、本地 K 线、财务指标和千问解读。"],
        ],
        col_widths=[2.8, 4.0, 8.4],
        caption="表5-1 前端页面入口",
    )

    add_h2(doc, "5.2 AI 选股使用说明")
    add_paragraph(
        doc,
        "进入 AI 选股页后，用户可以直接输入自然语言。对于普通问题，例如“这个 Agent 是什么”，系统返回普通对话；"
        "对于选股请求，例如“半导体行业市值 500 亿以上的龙头”，模型选择 stock_screen 工具，后端返回筛选条件和命中结果；"
        "对于“查看招商银行详情”，模型选择 stock_detail 工具并跳转或展示详情入口。"
    )
    add_table(
        doc,
        header=["示例输入", "期望行为"],
        rows=[
            ["查看招商银行详情", "定位 600036.SH 的股票详情，不执行筛选。"],
            ["银行股，PE低于15，PB低于1.2，股息率高于3%", "生成行业、PE、PB、股息率等结构化条件并执行筛选。"],
            ["帮我找 RPS 强势突破的股票", "选择 rps_breakout 策略并执行本地策略引擎。"],
            ["近三年净利润复合增速大于 20%", "识别为 unsupported metric，解释不支持原因，不筛选。"],
            ["你好", "作为普通对话回复，不展示工具轨迹。"],
        ],
        col_widths=[5.4, 9.8],
        caption="表5-2 AI 选股示例",
    )

    add_h2(doc, "5.3 条件选股与策略选股使用说明")
    add_paragraph(
        doc,
        "条件选股适合用户已经知道字段和阈值的场景。用户可选择行业、估值、财务和技术字段，设置大于、小于、区间、等于或包含，"
        "再点击执行。筛选结果支持排序、分页、批量选择和加入自选。保存后的条件策略可以在策略页选中后再执行，避免误触。"
    )
    add_paragraph(
        doc,
        "策略选股适合使用固定技术形态扫描，例如海龟突破、均线放量和 RPS 强势突破。策略结果不是投资评级，也不是历史回测，"
        "只是基于最新日线数据对当前命中股票进行排序。"
    )

    add_h2(doc, "5.4 自选股、预警和详情页使用说明")
    add_paragraph(
        doc,
        "用户登录后可以把筛选结果或详情页股票加入自选。自选明细支持按加入日期、涨跌幅、净资产收益率等字段排序；"
        "也可以批量编辑或批量移出。预警规则可以设置为价格突破、价格跌破、累计涨跌幅或日内涨跌幅，并可在页面直接启用、暂停或删除。"
    )
    add_paragraph(
        doc,
        "详情页优先展示本地数据库中的基本信息和 K 线。实时行情接口失败时，页面仍显示本地详情、最新已同步交易日和 K 线数据。"
        "千问解读按钮会调用 AI 生成分析文本，AI 不可用时只影响解读，不影响详情主数据。"
    )


def add_chapter_6(doc: Document):
    add_h1(doc, "6 系统测试")
    add_h2(doc, "6.1 测试环境与测试方法")
    add_paragraph(
        doc,
        "系统测试采用后端自动化测试、前端构建检查、浏览器 smoke、真实接口 smoke 和人工验收相结合的方法。"
        "普通 pytest 使用 fake model 或模拟响应，不依赖真实 AI；真实 Qwen/OpenAI-compatible 服务只在 release smoke 和"
        "agent reliability smoke 中验证。"
    )
    add_table(
        doc,
        header=["测试类别", "工具或命令", "覆盖内容"],
        rows=[
            ["后端单元/集成测试", "docker compose exec -T backend pytest", "认证、股票、筛选、Agent、策略、市场、通知、健康、缓存、调度器。"],
            ["发布冒烟", "python scripts/release_smoke.py", "Docker/HTTP 服务、AI health、数据健康、SSE、策略、详情和密钥扫描。"],
            ["Agent 真实 smoke", "python scripts/agent_reliability_smoke.py", "真实 query 的 tool、conditions、model_ms、tool_ms 和 fallback_reason。"],
            ["前端构建", "cd frontend && npm run build", "TypeScript/Vite 构建和生产包生成。"],
            ["前端 smoke", "smoke:auth/dashboard/strategy/chat/detail", "登录、行情、策略、对话、详情和路由状态恢复。"],
            ["人工验收", "真实浏览器", "验证码、登录、AI 选股、条件选股、策略、自选、详情和通知中心。"],
        ],
        col_widths=[3.2, 5.0, 7.0],
        caption="表6-1 测试方法",
    )

    add_h2(doc, "6.2 功能测试")
    add_table(
        doc,
        header=["编号", "测试内容", "预期结果", "结果"],
        rows=[
            ["T01", "验证码注册和登录", "验证码校验、注册成功、JWT 登录成功", "通过"],
            ["T02", "Dashboard 市场概览", "指数、板块、榜单和数据新鲜度正常展示", "通过"],
            ["T03", "AI 普通对话", "不调用筛选和策略工具，不展示工具轨迹", "通过"],
            ["T04", "AI 条件筛选", "模型选择 stock_screen，后端执行本地筛选", "通过"],
            ["T05", "AI 详情查询", "模型选择 stock_detail，不重新筛选", "通过"],
            ["T06", "unsupported metric", "不执行部分筛选，明确说明字段不支持", "通过"],
            ["T07", "条件选股", "多条件 AND/OR 筛选、排序、分页和批量加入自选", "通过"],
            ["T08", "策略选股", "6 个内置策略可执行，无数据时正常返回空结果", "通过"],
            ["T09", "自选与预警", "自选同步、排序、批量编辑和预警启停可用", "通过"],
            ["T10", "个股详情", "本地详情、K 线和千问解读入口可用", "通过"],
            ["T11", "通知中心", "触发预警后保存通知并支持已读", "通过"],
            ["T12", "数据健康", "展示 latest_trade_date、expected_trade_date、fresh、coverage 和 sync_warnings", "通过"],
        ],
        col_widths=[1.5, 4.4, 7.6, 1.7],
        caption="表6-2 功能测试结果",
    )

    add_h2(doc, "6.3 边界条件测试")
    add_table(
        doc,
        header=["边界场景", "控制规则", "验证结论"],
        rows=[
            ["AI 超时或上游不可达", "记录 model_ms 和 fallback_reason，不执行伪兜底筛选", "用户看到明确降级提示。"],
            ["不支持字段", "在 AI health、Qwen 和 screener 前停止", "不会返回只满足部分条件的结果。"],
            ["解释/排序/分页", "必须依赖上一轮上下文", "不会重新筛选。"],
            ["详情请求", "只执行 stock_detail", "不会误触发筛选。"],
            ["行业模糊词", "只展开到本地行业同义词", "不会把不存在的行业当作有效条件。"],
            ["K 线不足", "技术字段和策略跳过数据不足股票", "不合成行情。"],
            ["实时行情失败", "详情页使用本地数据兜底", "页面仍可查看。"],
            ["Redis 不可用", "缓存服务静默降级", "核心业务不受影响。"],
        ],
        col_widths=[4.4, 6.4, 4.4],
        caption="表6-3 边界条件测试",
    )

    add_h2(doc, "6.4 性能与数据健康测试")
    add_paragraph(
        doc,
        "性能测试重点关注 cold/warm 两类请求。Dashboard 聚合接口经过缓存预热后，市场概览、板块、异动榜和 ticker"
        "不应阻塞应用启动；/health/data 不允许伪造 fresh=true，只能在全市场日线覆盖到最近应有交易日时标记为新鲜。"
    )
    add_table(
        doc,
        header=["接口或任务", "检查项", "说明"],
        rows=[
            ["/api/v1/market/indices", "cold/warm 响应、6 个宽基指数", "首页市场概况使用。"],
            ["/api/v1/market/sectors?limit=20", "板块涨跌和行业数量", "板块列表按涨跌幅展示。"],
            ["/api/v1/market/movers?limit=10", "涨跌榜、成交额榜、换手率榜", "市场异动模块使用。"],
            ["/api/v1/market/ticker", "顶部滚动行情", "前端不应因失败阻塞主页面。"],
            ["/api/v1/health/data", "latest、expected、fresh、coverage、sync_warnings", "数据健康判断依据。"],
            ["strategy_select", "tool_ms 和缓存命中", "策略计算结果按最新交易日缓存。"],
            ["Agent SSE", "model_ms、tool_ms、fallback_reason", "定位模型慢路径和降级原因。"],
        ],
        col_widths=[5.0, 5.0, 5.2],
        caption="表6-4 性能与可观察性检查",
    )


def add_chapter_7(doc: Document):
    add_h1(doc, "7 总结与展望")
    add_paragraph(
        doc,
        "本项目完成了一个基于 FastAPI 和 Vue 的 A 股智能选股系统。系统不仅实现了传统的条件筛选，还引入了"
        "有边界的 AI Agent，使用户能够通过自然语言调用股票筛选、策略选股、结果解释、排序分页和详情查询等能力。"
        "后端通过字段白名单、schema 校验、unsupported metric 前置拦截和 fallback_reason 记录，降低了大模型输出不可控的风险。"
    )
    add_paragraph(
        doc,
        "数据层面，系统围绕股票基础信息、日线行情、财务指标、分红记录、自选股、预警通知和对话历史建立数据库模型，"
        "并通过定时任务和健康检查展示数据同步状态。前端方面，系统完成了登录、Dashboard、AI 选股、条件选股、策略选股、"
        "自选预警和详情页等主要页面，能够支撑课程演示和答辩展示。"
    )
    add_paragraph(
        doc,
        "后续可以从四个方向继续完善。第一，补充交易日历和更多数据源，减少节假日判断和深度数据缺失；第二，扩展财务时序数据，"
        "支持多期对比、三年 CAGR、现金流和 EPS 等字段；第三，将预警从前端轮询升级为后端任务或 WebSocket 推送；第四，"
        "在不牺牲可解释性的前提下增加策略回测和参数评估，使策略模块从当前扫描进一步扩展到历史验证。"
    )

    add_h1(doc, "参考文献")
    refs = [
        "[1] Ian Sommerville. Software Engineering[M]. 10th Edition. Pearson, 2015.",
        "[2] Grady Booch, James Rumbaugh, Ivar Jacobson. The Unified Modeling Language User Guide[M]. Addison-Wesley, 2005.",
        "[3] FastAPI. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com, 访问日期：2026-06-12.",
        "[4] SQLAlchemy. SQLAlchemy 2.0 Documentation[EB/OL]. https://docs.sqlalchemy.org/en/20/, 访问日期：2026-06-12.",
        "[5] Vue.js. Vue.js Guide[EB/OL]. https://vuejs.org, 访问日期：2026-06-12.",
        "[6] APScheduler. APScheduler Documentation[EB/OL]. https://apscheduler.readthedocs.io/, 访问日期：2026-06-12.",
        "[7] Baostock. Baostock 帮助文档[EB/OL]. http://baostock.com, 访问日期：2026-06-12.",
        "[8] AKShare. AKShare 开源金融数据接口库[EB/OL]. https://akshare.akfamily.xyz, 访问日期：2026-06-12.",
        "[9] 阿里云. 通义千问 DashScope API 文档[EB/OL]. https://help.aliyun.com/zh/dashscope/, 访问日期：2026-06-12.",
        "[10] OpenAI. Function Calling and Structured Outputs Documentation[EB/OL]. https://platform.openai.com/docs, 访问日期：2026-06-12.",
    ]
    for ref in refs:
        p = add_paragraph(doc, ref, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.left_indent = Pt(28)
        p.paragraph_format.first_line_indent = Pt(-28)


def build_document(out_path: Path):
    doc = Document()
    configure_document(doc)

    add_review_page(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_manual_toc(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_chapter_7(doc)

    # 防止最后一节沿用异常分节属性。
    doc.add_section(WD_SECTION.CONTINUOUS)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"文档已生成：{out_path}")


if __name__ == "__main__":
    output = Path(__file__).resolve().parent / "基于千问的股票筛选系统设计与实现.docx"
    build_document(output)
